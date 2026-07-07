"""
Dify Code Node — 报告模板填充与 Word 导出
============================================
适用：Dify 工作流中的 Code Node (Python 3)
输入变量：
  - template_file:   开始节点上传的 .docx 模板文件 (File)
  - llm_raw_output:  LLM 节点输出的 JSON 字符串 (String)
  - customer_name:   客户名称 (String)
  - eval_date:       评估日期 (String)
  - photo_file:      现场照片 (File, 可选)
输出变量：
  - word_output:     生成的 .docx 报告文件 (File)
"""

import json
import re
import os
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement


def main(template_file: str, llm_raw_output: str,
         customer_name: str, eval_date: str,
         photo_file: str = None) -> dict:
    # ============================================================
    # 1. 解析 LLM 输出
    # ============================================================
    try:
        llm_data = json.loads(llm_raw_output)
    except json.JSONDecodeError:
        json_match = re.search(r'\{.*\}', llm_raw_output, re.DOTALL)
        if json_match:
            llm_data = json.loads(json_match.group())
        else:
            raise ValueError("LLM 输出无法解析为 JSON")

    # ============================================================
    # 2. 构建扁平化的变量字典
    # ============================================================
    input_vars = {
        "input.customer_name": customer_name,
        "input.eval_date": eval_date,
    }

    llm_vars = {}
    for key, value in llm_data.items():
        llm_vars[f"llm_output.{key}"] = value

    all_vars = {**input_vars, **llm_vars}

    # ============================================================
    # 3. 加载模板并处理
    # ============================================================
    doc = Document(template_file)
    _process_all_paragraphs(doc, all_vars, photo_file)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_single_paragraph(para, all_vars)

    # ============================================================
    # 4. 保存输出
    # ============================================================
    output_path = '/tmp/generated_report.docx'
    doc.save(output_path)
    return {'word_output': output_path}


# ────────────────────────────────────────────────────────────────
#  辅助函数
# ────────────────────────────────────────────────────────────────

def _process_all_paragraphs(doc, all_vars, photo_file):
    """遍历文档所有段落，处理占位符"""
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 图片占位符 {{image:xxx}}
        if '{{image:' in para.text:
            _handle_image_placeholder(para, photo_file)
            continue

        # 普通占位符
        for key, value in all_vars.items():
            placeholder = '{{' + key + '}}'
            if placeholder not in para.text:
                continue

            if isinstance(value, list) and len(value) > 0:
                # 数组 → 在段落后插入 Word 表格
                _insert_table_after_paragraph(doc, para, value)
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, '')
            elif isinstance(value, str) and '\n' in value:
                # 含换行的文本 → 替换为多段落
                _replace_with_multi_paragraph(para, placeholder, value)
            else:
                # 普通文本/数字 → 直接替换
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))

    _cleanup_empty_paragraphs(doc)


def _process_single_paragraph(para, all_vars):
    """处理单个段落中的占位符替换（表格内单元格专用）"""
    for key, value in all_vars.items():
        placeholder = '{{' + key + '}}'
        if placeholder not in para.text:
            continue
        for run in para.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, str(value))


def _insert_table_after_paragraph(doc, para, data_list):
    """在段落后插入 Word 表格"""
    if not data_list or not isinstance(data_list, list):
        return

    headers = list(data_list[0].keys())
    parent = para._element.getparent()
    para_index = list(parent).index(para._element)

    rows_count = len(data_list) + 1  # +1 表头行
    cols_count = len(headers)
    table = doc.add_table(rows=rows_count, cols=cols_count)
    table.style = 'Light Grid Accent 1'

    # 表头
    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    # 数据行
    for row_idx, item in enumerate(data_list):
        for col_idx, header in enumerate(headers):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(item.get(header, ''))

    parent.insert(para_index + 1, table._element)


def _handle_image_placeholder(para, photo_file):
    """处理 {{image:xxx}} 图片占位符"""
    if not photo_file or not os.path.exists(photo_file):
        for run in para.runs:
            run.text = re.sub(r'\{\{image:\w+\}\}', '[图片未提供]', run.text)
        return

    for run in para.runs:
        run.text = ''

    run = para.add_run()
    try:
        run.add_picture(photo_file, width=Inches(5))
    except Exception:
        run.text = ''


def _replace_with_multi_paragraph(para, placeholder, text):
    """将含换行的文本替换为多个段落"""
    lines = text.split('\n')
    if not lines:
        return

    # 第一行替换原段落中的占位符
    for run in para.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, lines[0])

    # 后续行作为新段落插入
    if len(lines) > 1:
        parent = para._element.getparent()
        para_index = list(parent).index(para._element)

        for line in lines[1:]:
            new_para = OxmlElement('w:p')
            new_run = OxmlElement('w:r')
            new_text = OxmlElement('w:t')
            new_text.text = line
            new_run.append(new_text)
            new_para.append(new_run)
            para_index += 1
            parent.insert(para_index, new_para)


def _cleanup_empty_paragraphs(doc):
    """删除内容为空的段落"""
    for para in doc.paragraphs:
        if para.text.strip() == '' and len(para.runs) == 0:
            para._element.getparent().remove(para._element)
